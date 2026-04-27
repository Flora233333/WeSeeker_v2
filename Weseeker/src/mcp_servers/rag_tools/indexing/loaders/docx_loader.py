from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument
from langchain_core.documents import Document


def load_docx(path: Path) -> list[Document]:
    document = DocxDocument(path)
    sections: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        heading_level = _extract_heading_level(paragraph.style.name if paragraph.style else "")
        if heading_level is not None:
            sections.append(f"{'#' * heading_level} {text}")
        else:
            sections.append(text)

    for table in document.tables:
        table_markdown = _table_to_markdown(table)
        if table_markdown:
            sections.append(table_markdown)

    return [
        Document(
            page_content="\n\n".join(sections),
            metadata={
                "doc_type": "docx",
                "file_path": path.as_posix(),
                "file_name": path.name,
                "file_ext": path.suffix.lower(),
            },
        )
    ]


def _extract_heading_level(style_name: str) -> int | None:
    normalized = style_name.strip().lower()
    if not normalized.startswith("heading"):
        return None
    parts = normalized.split()
    if len(parts) < 2 or not parts[-1].isdigit():
        return None
    return max(1, min(int(parts[-1]), 6))


def _table_to_markdown(table) -> str:
    rows: list[list[str]] = []
    for row in table.rows:
        values = [cell.text.replace("\n", " ").strip() for cell in row.cells]
        if any(values):
            rows.append(values)

    if not rows:
        return ""

    header = rows[0]
    body = rows[1:] or [[]]
    header_line = "| " + " | ".join(header) + " |"
    separator_line = "| " + " | ".join(["---"] * len(header)) + " |"
    body_lines = ["| " + " | ".join(row) + " |" for row in body if row]
    return "\n".join([header_line, separator_line, *body_lines])
